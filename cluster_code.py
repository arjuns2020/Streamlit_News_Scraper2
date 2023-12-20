import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
from docx import Document
import spacy
from nltk.tokenize import word_tokenize

# Specify the full path to the spaCy model
model_path = r'C:\Users\arjun.murthy\Anaconda3\envs\autoviml\lib\site-packages\en_core_web_lg\en_core_web_lg-3.3.0'
nlp = spacy.load(model_path)


# Read Excel data into a DataFrame
df = pd.read_excel('oracle_datacentre_2023-11-01_2023-12-30_filtered.xlsx')

# Sample data loading
data = df
stop_words_list = list(ENGLISH_STOP_WORDS)
stop_words = list(ENGLISH_STOP_WORDS)

# Data preprocessing
df['text'] = df['text'].fillna('')
data = data[~data['text'].str.contains("Save my User ID and Password")]
data = data.dropna(subset=['text'])
data['text'] = data['text'].apply(lambda x: ' '.join([word.lower() for word in word_tokenize(x) if word.isalpha() and word.lower() not in stop_words]))

# Text vectorization using TF-IDF
vectorizer = TfidfVectorizer(max_features=1000, stop_words=stop_words_list)
X = vectorizer.fit_transform(df['text'])

# Clustering using K-Means
num_clusters = 5
kmeans = KMeans(n_clusters=num_clusters, random_state=42)
df['cluster'] = kmeans.fit_predict(X)

# Extractive summarization for each cluster using sumy
cluster_summaries = []
cluster_keywords = []

for i in range(num_clusters):
    cluster_articles = df[df['cluster'] == i]

    if not cluster_articles.empty:
        cluster_text = ' '.join(cluster_articles['text'])

        # Using LSA (Latent Semantic Analysis) Summarizer
        parser = PlaintextParser.from_string(cluster_text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        sentences_count = 3
        cluster_summary = ' '.join(str(sentence) for sentence in summarizer(parser.document, sentences_count))

        # Get keywords using TF-IDF
        terms = vectorizer.get_feature_names_out()
        cluster_tfidf_values = X[cluster_articles.index].toarray()
        avg_tfidf_scores = cluster_tfidf_values.mean(axis=0)
        top_keywords_idx = avg_tfidf_scores.argsort()[-5:][::-1]
        cluster_top_keywords = [terms[idx] for idx in top_keywords_idx]

        # Append the summary and keywords for each cluster
        cluster_summaries.append(cluster_summary)
        cluster_keywords.append(cluster_top_keywords)

# Print cluster information
for i, (summary, keywords) in enumerate(zip(cluster_summaries, cluster_keywords), 1):
    print(f"\nCluster {i} Summary:")
    print(f"Cluster Summary: {summary}")
    print(f"Cluster Keywords: {', '.join(keywords)}")

# Create a Word document and add content
doc = Document()
for i, (summary, keywords) in enumerate(zip(cluster_summaries, cluster_keywords), 1):
    doc.add_heading(f'Cluster {i} Summary', level=2)
    doc.add_paragraph(summary)
    doc.add_heading(f'Cluster {i} Keywords', level=2)
    doc.add_paragraph(', '.join(keywords))

# Save the document
doc.save('summary_document.docx')
