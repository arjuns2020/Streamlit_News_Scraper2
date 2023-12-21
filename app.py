import os
import pandas as pd
from gnews import GNews
from datetime import datetime
from langid import classify
from joblib import Parallel, delayed
import streamlit as st
import xlsxwriter
from io import BytesIO
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
from docx import Document
import spacy
from nltk.tokenize import word_tokenize
#st.set_option('client.showWarningOnDirectSt.textWrites', False)

@st.cache_resource
def download_en_core_web_sm():
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    
nlp = spacy.load("en_core_web_sm")
nltk.download('punkt')


class Session:
    def __init__(self):
        self.query = None
        self.start_date = None
        self.end_date = None
        self.google_news = None
        self.results_list = None
        self.df = None
        self.display_results = False  # Store display state within session state
        

# Initialize Streamlit session state
if "results_list" not in st.session_state:
    st.session_state.results_list = None

if "df" not in st.session_state:
    st.session_state.df = None

if "display_results" not in st.session_state:
    st.session_state.display_results = False



# Function to check if the article is in English
def is_english(text):
    lang, _ = classify(text)
    return lang == 'en'

# Function to get full article using newspaper3k and perform NLP
@st.cache_data()
def get_full_article_with_nlp(url):
    article = GNews().get_full_article(url)

    # Perform NLP on the article
    try:
        article.download()
        article.parse()
        article.nlp()
    except Exception as e:
        print(f"Error performing NLP on the article: {e}")

    return {
        'title': article.title,
        'text': article.text,
        'authors': article.authors,
        'summary': article.summary,
        'keywords': article.keywords,
    }
    
@st.cache_data(persist=True)
def process_article(article):
    source = article.get('source', 'N/A')
    if source == 'N/A':
        # Check if the source is still not available and try to extract it from the URL
        url = article.get('url', '')
        if url:
            source_from_url = url.split('/')[2]  # Extract domain from the URL
            source = source_from_url if source_from_url else 'N/A'

    if is_english(article.get('title', '')):
        title = article.get('title', 'N/A')

        # Extract information after " ... - " or " - " in the title
        if ' ... - ' in title:
            title_parts = title.split(' ... - ')
        elif ' - ' in title:
            title_parts = title.split(' - ')
        else:
            title_parts = [title]

        article_source = title_parts[1].strip() if len(title_parts) > 1 else 'N/A'

        result = {
            'Title': article.get('title', 'N/A'),
            'Source': source,
            'Article_Source': article_source,
            'URL': article.get('url', 'N/A'),
        }

        try:
            # Get the full article using newspaper3k and perform NLP
            full_article = get_full_article_with_nlp(article.get('url'))
            result.update(full_article)
        except Exception as e:
            print(f"Error processing article: {e}")

        return result
    else:
        return None

# Your clustering function
@st.cache_data(persist=True)
def run_clustering(df):
    # Sample data loading
    data = df
    stop_words_list = list(ENGLISH_STOP_WORDS)
    stop_words = list(ENGLISH_STOP_WORDS)

    # Data preprocessing
    df['text'] = df['text'].fillna('')
    data = data[~data['text'].str.contains(" User ID and Password")]
    data = data.dropna(subset=['text'])
    data['text'] = data['text'].apply(
        lambda x: ' '.join([word.lower() for word in word_tokenize(x) if word.isalpha() and word.lower() not in stop_words])
    )

    # Text vectorization using TF-IDF
    vectorizer = TfidfVectorizer(max_features=1000, stop_words=stop_words_list)
    X = vectorizer.fit_transform(df['text'])

    # Clustering using K-Means
    num_clusters = 5
    kmeans = KMeans(n_clusters=num_clusters, random_state=42)
    df['cluster'] = kmeans.fit_predict(X)

    # Extractive summarization for each cluster using sumy
    cluster_summaries = []
    cluster_keywords = []

    for i in range(num_clusters):
        cluster_articles = df[df['cluster'] == i]

        if not cluster_articles.empty:
            cluster_text = ' '.join(cluster_articles['text'])

            # Using LSA (Latent Semantic Analysis) Summarizer
            parser = PlaintextParser.from_string(cluster_text, Tokenizer("english"))
            summarizer = LsaSummarizer()
            sentences_count = 3
            cluster_summary = ' '.join(
                str(sentence) for sentence in summarizer(parser.document, sentences_count)
            )

            # Get keywords using TF-IDF
            terms = vectorizer.get_feature_names_out()

            # Validate indices to avoid IndexError
            valid_indices = [idx for idx in cluster_articles.index if idx < X.shape[0]]
            cluster_tfidf_values = X[valid_indices].toarray()

            avg_tfidf_scores = cluster_tfidf_values.mean(axis=0)
            top_keywords_idx = avg_tfidf_scores.argsort()[-5:][::-1]
            cluster_top_keywords = [terms[idx] for idx in top_keywords_idx]

            # Append the summary and keywords for each cluster
            cluster_summaries.append(cluster_summary)
            cluster_keywords.append(cluster_top_keywords)

    return cluster_summaries, cluster_keywords

def main():
    st.title("News Scraper, Analyzer, and Clusterer")
    st.sidebar.header("Search Parameters")

    # User input
    query = st.sidebar.text_input("Enter Keywords (separated by space)", 'aws cloud')
    start_date = st.sidebar.date_input("Enter Start Date", datetime(2023, 11, 1))
    end_date = st.sidebar.date_input("Enter End Date", datetime(2023, 12, 30))

    # Set the language filter
    GNews.language = 'en'

    # Initialize GNews object with query parameters
    google_news = GNews(country='US', max_results=100)

    # Search button
    if st.sidebar.button("Search"):
        with st.spinner("Please wait while scraping the web..."):
            # Set the start and end dates based on UI input
            google_news.start_date = start_date
            google_news.end_date = end_date

            # Get news results
            news_results = google_news.get_news(query)

            # Parallel processing using joblib
            results_list = Parallel(n_jobs=-1)(delayed(process_article)(article) for article in news_results)

            # Filter out None values (articles not in English)
            results_list = [result for result in results_list if result is not None]

            # Create a DataFrame from the list of results
            df = pd.DataFrame(results_list)

            # Filter out null values (articles not in English or with null title/text)
            df = df.dropna(subset=['Title', 'text'])

            # Display results
            st.subheader("Results")
            st.dataframe(df)

            # Initialize cluster_summaries and cluster_keywords in case of an exception
            cluster_summaries, cluster_keywords = [], []

            # Call the clustering function
            try:
                cluster_summaries, cluster_keywords = run_clustering(df)
            except Exception as e:
                st.error(f"An error occurred during clustering: {str(e)}")

            # Display cluster information
            st.subheader("Cluster Information")

            # Initialize an empty string to store cluster information
            cluster_info = ""

            # Iterate over clusters and concatenate information into a single string
            for i, (summary, keywords) in enumerate(zip(cluster_summaries, cluster_keywords), 1):
                cluster_info += f"Cluster {i} Summary: {summary}\n"
                cluster_info += f"Cluster {i} Keywords: {', '.join(keywords)}\n\n"

            # Display the concatenated information in a larger text area with scrolling
            st.text_area("Cluster Information", value=cluster_info, height=500, key="cluster_info")

            # Download button for cluster results
            if st.button("Download Cluster Results"):
                # Save the document
                doc = Document()
                for i, (summary, keywords) in enumerate(zip(cluster_summaries, cluster_keywords), 1):
                    doc.add_heading(f'Cluster {i} Summary', level=2)
                    doc.add_paragraph(summary)
                    doc.add_heading(f'Cluster {i} Keywords', level=2)
                    doc.add_paragraph(', '.join(keywords))

                # Save the document
                doc.save('cluster_summary.docx')

            # Store results in session state
            st.session_state.results_list = results_list
            st.session_state.df = df
            st.session_state.display_results = True

    if st.sidebar.button("Refresh"):
        # Clear the results and reset the UI
        st.session_state.results_list = None
        st.session_state.df = None
        st.session_state.display_results = False
        st.experimental_rerun()

    # Display results only if the display flag is True
    if st.session_state.display_results:
        st.subheader("Results")
        st.dataframe(st.session_state.df)

    # Download buttons in the sidebar
    if st.session_state.results_list:
        # Download button for the filtered DataFrame
        excel_file_path_filtered = (
            f"{query.split()[0]}_{query.split()[1]}_filtered.xlsx"
        )

        # Save the filtered DataFrame to an Excel file
        st.session_state.df.to_excel(excel_file_path_filtered, index=False, engine='openpyxl')

        # Save the filtered DataFrame to the archive folder with the current date
        archive_folder = os.path.join("archive", datetime.now().strftime("%Y-%m-%d"))
        os.makedirs(archive_folder, exist_ok=True)
        excel_file_path_archive = os.path.join(archive_folder, excel_file_path_filtered.split(os.path.sep)[-1])
        st.session_state.df.to_excel(excel_file_path_archive, index=False, engine='openpyxl')

        # Create an in-memory Excel file for the filtered DataFrame
        output_filtered = BytesIO()
        st.session_state.df.to_excel(output_filtered, index=False, engine='openpyxl')

        # Download button for the filtered DataFrame in the sidebar
        st.sidebar.download_button(
            label="Download Filtered Results",
            data=output_filtered.getvalue(),
            file_name=excel_file_path_filtered,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="filtered_excel",
            help="Click to download the filtered results",
        )

        # Save the document for clustering results
        doc = Document()
        for i, (summary, keywords) in enumerate(zip(cluster_summaries, cluster_keywords), 1):
            doc.add_heading(f'Cluster {i} Summary', level=2)
            doc.add_paragraph(summary)
            doc.add_heading(f'Cluster {i} Keywords', level=2)
            doc.add_paragraph(', '.join(keywords))
            
        # Save the document to a BytesIO object
        output_doc = BytesIO()
        doc.save(output_doc)
        output_doc.seek(0)  # Move the cursor to the beginning of the BytesIO object

        # Download button for the cluster summary document in the sidebar
        st.sidebar.download_button(
            label="Download Cluster Summary",
            data=output_doc.getvalue(),
            file_name=f"{query.split()[0]}_{query.split()[1]}_cluster_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="cluster_summary_docx",
            help="Click to download the cluster summary document",
        )

if __name__ == "__main__":
    main()
